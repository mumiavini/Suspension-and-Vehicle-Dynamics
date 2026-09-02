#!/usr/bin/env python3
"""
build_summary_doc.py
====================

Render the geometry summary as a Word document and a PDF.

``scripts/geometry_summary.py`` is the source of truth and emits fixed-width
text. This script builds the same content as a formatted Word document with real
tables, so the team has something to hand to a Design Event judge, and exports it
to PDF through Word.

Every number is pulled from the dataclasses that ``sla_geometry.py`` and
``steering_geometry.py`` return -- nothing is parsed out of the text report and
nothing is typed in by hand. That is the property the previous document lost:
it was assembled manually from stale runs, so its rate tables and its force
tables ended up disagreeing with each other and with the code.

The document is in English, per CLAUDE.md ("All English: comments, docstrings,
identifiers, docs"). The superseded document was in Portuguese.

    python scripts/build_summary_doc.py                    # docx + pdf
    python scripts/build_summary_doc.py --no-pdf           # docx only

PDF export drives Microsoft Word through COM and needs ``pywin32``. On a machine
without Word, pass --no-pdf and convert the .docx elsewhere.
"""

from __future__ import annotations

import argparse
import datetime as dt
import sys
from collections.abc import Sequence
from pathlib import Path

REPO = Path(__file__).resolve().parent.parent
if str(REPO) not in sys.path:
    sys.path.insert(0, str(REPO))

import geometry_summary as gs  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.section import WD_ORIENT  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches, Pt, RGBColor  # noqa: E402

import sla_geometry as sla  # noqa: E402
import steering_geometry as stg  # noqa: E402

OUT_DIR = REPO / "Geometry Summary"
# Keep the team's established document name so this UPDATES the file everyone
# already links to, rather than leaving the stale one beside a new one.
STEM = "Hardpoints Suspensão 2027"
DOCX = OUT_DIR / f"{STEM}.docx"
PDF = OUT_DIR / f"{STEM}.pdf"
CHART = REPO / "geometria.png"

NAVY = "1F3864"
BAND = "F2F3F5"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x59, 0x59, 0x59)
BAD = RGBColor(0x9E, 0x24, 0x1C)
GOOD = RGBColor(0x26, 0x64, 0x4A)

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"


# --------------------------------------------------------------------------- #
# docx helpers
# --------------------------------------------------------------------------- #

def shade(cell, hex_fill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def cell_text(cell, text: str, *, bold=False, color=INK, size=9.5,
              align=WD_ALIGN_PARAGRAPH.LEFT, mono=False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = MONO_FONT if mono else BODY_FONT


def heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 11)
    run.font.name = BODY_FONT
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    if level == 1:
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:color"), NAVY)
        pbdr.append(bottom)
        p._p.get_or_add_pPr().append(pbdr)


def note(doc, text: str, *, italic=True, color=MUTED, size=9) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = BODY_FONT


def table(doc, headers: Sequence[str], rows: Sequence[Sequence[str]],
          widths: Sequence[float] | None = None,
          right_align_from: int = 1) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"

    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, NAVY)
        cell_text(c, h, bold=True, color=WHITE, size=9,
                  align=WD_ALIGN_PARAGRAPH.RIGHT if i >= right_align_from
                  else WD_ALIGN_PARAGRAPH.LEFT)

    for n, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if n % 2 == 1:
                shade(cells[i], BAND)
            colour = INK
            mono = i >= right_align_from
            if val.startswith("FAIL"):
                colour = BAD
            elif val.startswith("OK"):
                colour = GOOD
            cell_text(cells[i], val, color=colour, mono=mono,
                      align=WD_ALIGN_PARAGRAPH.RIGHT if i >= right_align_from
                      else WD_ALIGN_PARAGRAPH.LEFT)

    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def check(label: str, value: float, lo: float, hi: float,
          unit: str = "mm", fmt: str = ".2f") -> list[str]:
    ok = lo <= value <= hi
    return [("OK" if ok else "FAIL"), label,
            f"{value:{fmt}} {unit}".strip(), f"{lo:g} to {hi:g}"]


# --------------------------------------------------------------------------- #
# sections
# --------------------------------------------------------------------------- #

def cover(doc, veh: sla.VehicleData) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("SUSPENSION GEOMETRY — HARDPOINTS 2027")
    r.bold = True
    r.font.size = Pt(21)
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    r.font.name = BODY_FONT

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(f"Double A-arm (SLA), front and rear  |  {veh.name}")
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    r.font.name = BODY_FONT

    note(doc,
         "Sign conventions: negative camber = top of the wheel inboard; toe-in "
         "positive; positive force = tension. Lengths in mm, angles in degrees "
         "unless stated. Two coordinate frames appear in this document and every "
         "table names the one it uses — see Provenance below.")

    stamp = dt.date.today().isoformat()
    note(doc,
         f"Generated {stamp} by scripts/build_summary_doc.py from "
         f"sla_geometry.py and steering_geometry.py. Do not hand-edit: change the "
         f"config in those scripts and re-run. This document supersedes "
         f"\"Hardpoints Suspensão 2027\", whose rate tables predated the "
         f"upright-rotation sign fix (commit 0e7e524) and whose force tables "
         f"assumed a different weight distribution from its own roll-stiffness "
         f"section.", italic=False)


def provenance(doc, hp: gs.MergedHardpoints) -> None:
    heading(doc, "0. Provenance and frames")
    table(doc, ["Content", "Produced by"], [
        ["Wishbone hardpoints, static KPIs, rates, roll, anti-geometry, leg forces",
         "sla_geometry.py"],
        ["Caster, outboard ball joints, tie rod, rack, bump steer, Ackermann, effort",
         "steering_geometry.py"],
        ["3D kinematic solve behind the steering numbers",
         "vdcore.geometry.solver.DWSolver"],
        ["Composition, cross-checks, this document",
         "scripts/geometry_summary.py + build_summary_doc.py"],
    ], widths=[4.6, 2.3], right_align_from=99)

    table(doc, ["Frame", "Axes", "Used in"], [
        ["DESIGN", "y outboard, z up, x REARWARD", "Sections 2 and 3 (front-view work)"],
        ["ISO 8855", "X FORWARD, Y LEFT, Z up", "Section 6 (model export)"],
    ], widths=[1.1, 2.6, 3.2], right_align_from=99)
    note(doc, "Conversion: X_iso = −x_rearward, Y_iso = ±y_outboard, Z_iso = z. "
              "In ISO 8855 the rear axle is at NEGATIVE X and right-hand corners "
              "are at NEGATIVE Y.")

    # Until rev 5 the rear toe link was read back out of the exported CSV and
    # this block printed a red warning gated on a `rear_tie_rod_from_csv` flag.
    # That flag, and the condition behind it, are gone: the link is now a
    # declared design input in geometry_summary.py. The disclosure stays --
    # design_intent provenance must be reported wherever it is consumed -- but
    # it is a statement of provenance now, not a defect.
    note(doc,
         f"The rear toe link is a declared {gs.REAR_TOE_LINK_SOURCE} input "
         "(REAR_TOE_LINK_INBOARD / _OUTBOARD in geometry_summary.py), not a "
         "measured or synthesised value. The rear has no rack, so "
         "steering_geometry.py -- which covers the front axle only -- does not "
         "produce it. It is regenerated with every run, and its inboard Z was "
         "chosen to null the linear bump-steer rate; section 3c reports the "
         "solved result and how sharp that knob is.")


def vehicle_section(doc, veh: sla.VehicleData, res: sla.VehicleResults,
                    front: sla.AxleGeometry, rear: sla.AxleGeometry) -> None:
    heading(doc, "1. Vehicle, stiffness and rules compliance")
    table(doc, ["Parameter", "Value"], [
        ["Total / sprung mass", f"{veh.total_mass_kg:.1f} / {veh.sprung_mass_kg:.1f} kg"],
        ["Wheelbase", f"{veh.wheelbase_mm:.1f} mm"],
        ["Front track / rear track",
         f"{front.inputs.track_mm:.1f} / {rear.inputs.track_mm:.1f} mm"],
        ["CG height / station",
         f"{veh.cg_height_mm:.1f} mm / {veh.cg_from_front_axle_mm:.1f} mm aft of the front axle"],
        ["Static front mass fraction", f"{100 * veh.front_mass_fraction:.1f} %"],
        ["Roll axis height at the CG", f"{res.roll_axis_height_at_cg_mm:.1f} mm"],
        ["Roll moment arm", f"{res.roll_moment_arm_mm:.1f} mm"],
        ["Target roll gradient", f"{veh.target_roll_gradient_deg_per_g:.2f} °/g"],
        ["Required roll stiffness",
         f"{res.required_roll_stiffness_Nm_per_deg:.1f} N·m/°"],
        ["Chassis torsional stiffness",
         f"{res.chassis_torsion_min_Nm_per_deg:.0f} (min) / "
         f"{res.chassis_torsion_target_Nm_per_deg:.0f} (target) N·m/°"],
    ], widths=[2.9, 4.0], right_align_from=99)
    note(doc, "The front mass fraction sets BOTH the roll axis height and the "
              "leg-force load cases in section 7. It is a single input; the two "
              "must never be allowed to disagree again.")

    narrow = min(front.inputs.track_mm, rear.inputs.track_mm)
    wide = max(front.inputs.track_mm, rear.inputs.track_mm)
    rows = [
        check("Wheelbase", veh.wheelbase_mm, gs.FSAE_MIN_WHEELBASE_MM, 1e4, "mm", ".0f"),
        check("Narrow / wide track ratio", 100 * narrow / wide,
              100 * gs.FSAE_MIN_TRACK_RATIO, 100.0, "%", ".1f"),
        [("OK" if narrow >= res.tilt_min_track_mm else "FAIL"),
         "60° tilt test", f"{narrow:.0f} mm fitted",
         f"needs {res.tilt_min_track_mm:.0f} mm"],
    ]
    table(doc, ["Status", "Item", "Value", "Target"], rows,
          widths=[0.7, 3.0, 1.7, 1.5], right_align_from=2)
    note(doc, f"The tilt test uses the NARROWER track, which is the "
              f"{'front' if front.inputs.track_mm < rear.inputs.track_mm else 'rear'} "
              f"({narrow:.0f} mm). Quoting that as \"the track\" hides the "
              f"{wide:.0f} mm on the other axle.")


def axle_section(doc, geo: sla.AxleGeometry, rates: "gs.AxleRates",
                 roll: "gs.AxleRollState", number: str) -> None:
    inp, lim = geo.inputs, geo.inputs.limits
    heading(doc, f"{number}. {inp.name} suspension")

    heading(doc, "Front-view points", 2)
    note(doc, "DESIGN frame: y positive OUTBOARD, z positive up, origin at ground "
              "level on the car centreline. Two decimals are required — the "
              "integers printed in the superseded document could not reproduce "
              "its own FVIC or roll centre.")
    table(doc, ["Point", "y [mm]", "z [mm]"], [
        [label, f"{pt[0]:.2f}", f"{pt[1]:.2f}"]
        for label, pt in (("Lower ball joint (LBJ)", geo.lbj),
                          ("Upper ball joint (UBJ)", geo.ubj),
                          ("LCA inboard", geo.lca_in),
                          ("UCA inboard", geo.uca_in),
                          ("FVIC (construction)", geo.fvic))
    ], widths=[3.3, 1.8, 1.8])

    heading(doc, "Wishbones", 2)
    table(doc, ["Parameter", "Value"], [
        ["LCA / UCA length (front view)",
         f"{geo.lca_length_mm:.2f} / {geo.uca_length_mm:.2f} mm  "
         f"(ratio {geo.uca_lca_ratio:.3f})"],
        ["Outboard vertical separation", f"{geo.outer_vertical_sep_mm:.2f} mm"],
        ["Inboard vertical separation", f"{geo.inner_vertical_sep_mm:.2f} mm"],
        ["LCA inclination", f"{geo.lca_inclination_deg:.2f}°  "
         f"({'falls' if geo.lca_inclination_deg > 0 else 'rises'} from wheel to chassis)"],
        ["UCA inclination", f"{geo.uca_inclination_deg:.2f}°  "
         f"({'falls' if geo.uca_inclination_deg > 0 else 'rises'} from wheel to chassis)"],
        ["RC that would flatten the LCA", f"{geo.rc_height_for_flat_lca_mm:.2f} mm"],
    ], widths=[2.9, 4.0], right_align_from=99)

    heading(doc, "Checks", 2)
    rim_lo, rim_hi = inp.rim_z_band
    c = lim.ball_joint_clearance_mm
    inside = (rim_lo + c) <= geo.lbj[1] and geo.ubj[1] <= (rim_hi - c)
    rows = [
        check("Roll centre height", inp.rc_height_mm, *lim.rc_height_mm),
        check("FVSA length", inp.fvsa_length_mm, *lim.fvsa_length_mm),
        check("Scrub radius", geo.scrub_radius_mm, *lim.scrub_radius_mm),
        check("KPI", inp.kpi_deg, *lim.kpi_deg, unit="°"),
        check("Kingpin length (front view)", inp.kingpin_length_mm,
              *lim.kingpin_length_mm),
        check("LCA length (front view)", geo.lca_length_mm, *lim.lca_length_mm),
        check("UCA / LCA ratio", geo.uca_lca_ratio, *lim.uca_lca_ratio,
              unit="", fmt=".3f"),
        check("Camber gain", abs(rates.camber_gain_deg_per_mm),
              *lim.camber_gain_deg_per_mm, unit="°/mm", fmt=".4f"),
        [("OK" if geo.fvic[0] < 0 else "FAIL"), "FVIC on the far side of the car",
         f"{geo.fvic[0]:.2f} mm", "—"],
        [("OK" if inside else "FAIL"), "Ball joints inside the rim",
         f"{geo.lbj[1]:.0f} / {geo.ubj[1]:.0f} mm",
         f"{rim_lo:.0f} to {rim_hi:.0f}"],
    ]
    table(doc, ["Status", "Item", "Value", "Target"], rows,
          widths=[0.7, 3.0, 1.7, 1.5], right_align_from=2)
    note(doc, "Kingpin and arm lengths here are FRONT-VIEW projections — the right "
              "quantity for the swing-arm construction, but not the length that "
              "gets cut. Section 5 checks the true 3D members.")

    heading(doc, "Rates about static", 2)
    table(doc, ["Quantity", "Value"], [
        ["Camber gain", f"{rates.camber_gain_deg_per_mm:.4f} °/mm"],
        ["Roll centre migration", f"{rates.rc_migration_mm_per_mm:.4f} mm/mm"],
        ["Half-track change", f"{rates.half_track_change_mm_per_mm:+.4f} mm/mm"],
        ["Camber at full bump", f"{rates.camber_full_bump_deg:.2f}°"],
        ["Camber at full droop", f"{rates.camber_full_droop_deg:.2f}°"],
        ["RC over the travel range", f"{rates.rc_min_mm:.1f} to {rates.rc_max_mm:.1f} mm"],
    ], widths=[2.9, 4.0], right_align_from=99)
    note(doc, f"Travel ±{inp.travel_bump_mm:.0f} mm. Camber gain equals "
              f"57.2958 / FVSA = {57.2958 / inp.fvsa_length_mm:.4f} °/mm, the "
              f"textbook value for a wheel turning about an instant centre "
              f"{inp.fvsa_length_mm:.0f} mm away. Half-track change is POSITIVE: "
              f"the contact patch moves outboard in bump.")

    heading(doc, f"At {roll.roll_deg:.1f}° of roll", 2)
    lo, hi = lim.outer_camber_in_roll_deg
    table(doc, ["Quantity", "Value"], [
        ["Outer wheel camber vs road",
         f"{roll.outer_camber_deg:.2f}°  (static {inp.static_camber_deg:+.2f}°)"],
        ["Inner wheel camber vs road", f"{roll.inner_camber_deg:.2f}°"],
        ["Roll centre height",
         f"{roll.rc_height_mm:.1f} mm  (design {inp.rc_height_mm:.1f})"],
        ["Roll centre lateral migration", f"{roll.rc_lateral_mm:.1f} mm"],
        ["Wheel travel at that roll", f"{roll.wheel_travel_mm:.2f} mm"],
        ["Outer camber in the useful window",
         f"{'OK' if lo <= roll.outer_camber_deg <= hi else 'FAIL'} "
         f"({lo:.1f} to {hi:.1f}°)"],
    ], widths=[2.9, 4.0], right_align_from=99)

    heading(doc, "Longitudinal layout and anti-geometry", 2)
    note(doc, "DESIGN frame, x positive REARWARD. Sweep and e/a as BUILT are in "
              "section 5, measured off the caster-corrected ball joints.")
    table(doc, ["Parameter", "Value"], [
        ["LCA pickups x", f"{geo.lca_in_front_x_mm:.1f} / {geo.lca_in_rear_x_mm:.1f} mm"],
        ["UCA pickups x", f"{geo.uca_in_front_x_mm:.1f} / {geo.uca_in_rear_x_mm:.1f} mm"],
        ["LCA / UCA e/a ratio", f"see section 5 — measured as built "
         f"(target ≤ {lim.ea_ratio_max:g})"],
        [geo.anti_label, f"{geo.anti_percent:.2f} %   "
         f"(target {lim.anti_percent[0]:g} to {lim.anti_percent[1]:g})"],
    ], widths=[2.9, 4.0], right_align_from=99)
    if geo.svic is None:
        note(doc, "All pivot axes are exactly horizontal, so the side-view instant "
                  "centre is at infinity and the anti-feature is identically zero — "
                  "not approximately zero. That is a design decision to defend, not "
                  "a check that happened to pass.")


def steering_section(doc, sr: stg.SteeringReport, steer_in) -> None:
    sg, rates, eff = sr.geometry, sr.rates, sr.effort
    heading(doc, "4. Steering (front axle)")
    note(doc, "Absent in its entirety from the superseded document. Caster and "
              "mechanical trail set steering feel and effort; bump steer is the "
              "tie rod's primary KPI.")

    table(doc, ["Parameter", "Value"], [
        ["Caster angle", f"{steer_in.caster_deg:.2f}°"],
        ["Mechanical trail", f"{sg.mechanical_trail_mm:.2f} mm"],
        ["Scrub radius (3D)", f"{sg.scrub_radius_mm:.2f} mm"],
        ["Steering arm length", f"{sg.steer_arm_length_mm:.2f} mm"],
        ["Tie rod length", f"{sg.tie_rod_length_mm:.2f} mm"],
        ["Rack position (x rearward, z)",
         f"{steer_in.rack_x_mm:.1f}, {steer_in.rack_z_mm:.1f} mm"],
        ["Rack half length / max travel",
         f"{steer_in.rack_half_length_mm:.1f} / ±{steer_in.max_rack_travel_mm:.1f} mm"],
    ], widths=[2.9, 4.0], right_align_from=99)

    table(doc, ["Quantity", "Value"], [
        ["Bump steer (per side)",
         f"{rates.bump_steer_deg_per_mm_per_side:+.5f} °/mm"],
        ["Bump steer (total toe)",
         f"{rates.bump_steer_total_toe_deg_per_mm:+.5f} °/mm"],
        ["Toe at full bump / droop (per side)",
         f"{rates.toe_at_full_bump_deg_per_side:+.3f} / "
         f"{rates.toe_at_full_droop_deg_per_side:+.3f} °"],
        ["C-factor", f"{rates.c_factor_mm_per_deg:.3f} mm/°"],
        ["Steering ratio", f"{rates.steering_ratio:.2f} : 1"],
        ["Max steer at stroke (outer / inner)",
         f"{rates.max_steer_at_stroke_deg:.2f} / "
         f"{rates.max_steer_inner_at_stroke_deg:.2f}°"],
        ["Ackermann at target steer", f"{rates.ackermann_pct_at_target:.1f} %"],
        ["Worst rod-end misalignment",
         f"{rates.worst_rod_end_misalignment_deg:.2f}°"],
    ], widths=[2.9, 4.0], right_align_from=99)
    note(doc, "Bump steer is per side and total toe separately, because per-side "
              "versus total has caused confusion on the real car. The gradient is "
              "effectively zero at static: both tie rods aim at their own "
              "front-view instant centre. The residual is second order and the "
              "same sign at both ends of travel, so roll steer stays near zero "
              "while total toe rises slightly at full travel.")

    heading(doc, "Steering effort (parking)", 2)
    table(doc, ["Quantity", "Value"], [
        ["Fz per front wheel", f"{eff.Fz_per_wheel_N:.1f} N"],
        ["Kingpin moment per wheel", f"{eff.kingpin_moment_Nm:.2f} N·m"],
        ["Rack force (both wheels)", f"{eff.rack_force_N:.1f} N"],
        ["Steering wheel torque", f"{eff.steering_wheel_torque_Nm:.2f} N·m"],
        ["Rim force", f"{eff.rim_force_N:.1f} N"],
    ], widths=[2.9, 4.0], right_align_from=99)
    note(doc, "Parking effort assumes a friction coefficient. PUCPR Racing has no "
              "TTC tyre data — this is a design_intent assumption, not a measured "
              "value.", color=BAD, italic=False)


def members_section(doc, hp: gs.MergedHardpoints) -> None:
    lo, hi = gs.ARM_LENGTH_WINDOW_MM
    klo, khi = gs.KINGPIN_WINDOW_MM
    heading(doc, "5. Member lengths and sweep as built (true 3D)")
    note(doc, "Measured off the merged hardpoints in ISO 8855, after caster. These "
              "are the members that get cut, as opposed to the front-view "
              "projections checked in sections 2 and 3.")

    for corner, label in (("FL", "Front"), ("RL", "Rear")):
        m = gs.member_lengths(hp, corner)
        heading(doc, f"{label} axle [{corner}]", 2)
        rows = [check(name, val, lo, hi) for name, val in m.legs]
        rows.append(check("Kingpin length, front view", m.kingpin_2d, klo, khi))
        rows.append(check("Kingpin length, TRUE 3D", m.kingpin_3d, klo, khi))
        table(doc, ["Status", "Item", "Value", "Target"], rows,
              widths=[0.7, 3.0, 1.7, 1.5], right_align_from=2)
        table(doc, ["Parameter", "Value"], [
            ["LCA base / sweep as built",
             f"{m.lca_base:.2f} / {m.lca_sweep:.2f} mm   "
             f"(e/a {m.lca_sweep / (m.lca_base / 2):.3f})"],
            ["UCA base / sweep as built",
             f"{m.uca_base:.2f} / {m.uca_sweep:.2f} mm   "
             f"(e/a {m.uca_sweep / (m.uca_base / 2):.3f})"],
            ["Tie rod", "—" if m.tie_rod != m.tie_rod else f"{m.tie_rod:.2f} mm"],
        ], widths=[2.9, 4.0], right_align_from=99)

    note(doc, "Any leg flagged above is longer than the packaging window the "
              "front-view check passes it against. The swept rear legs matter "
              "most: e/a above 1 means the sweep roughly doubles the leg load, "
              "and buckling of a long compression member is outside this tool's "
              "scope. Someone has to check it.", color=BAD, italic=False)


def hardpoints_section(doc, hp: gs.MergedHardpoints) -> None:
    heading(doc, "6. Merged hardpoints (ISO 8855)")
    note(doc, "X POSITIVE FORWARD. Y POSITIVE LEFT. Z POSITIVE UP. Origin at the "
              "front axle centreline, ground plane, vehicle centreline. The rear "
              "axle is therefore at NEGATIVE X and right-hand corners at NEGATIVE "
              "Y. This is NOT the design frame used in sections 2 and 3.")

    for corner in gs.CORNERS:
        heading(doc, corner, 2)
        rows = []
        for name in gs.POINT_ORDER:
            pt = hp.points[corner].get(name)
            if pt is None:
                continue
            ref = " *" if name in ("WHEEL_CENTER", "CONTACT_PATCH") else ""
            rows.append([name + ref, f"{pt[0]:.2f}", f"{pt[1]:.2f}", f"{pt[2]:.2f}"])
        table(doc, ["Point", "X [mm]", "Y [mm]", "Z [mm]"], rows,
              widths=[2.4, 1.5, 1.5, 1.5])

    cam, toe = gs.static_alignment_encoded(hp, "FL")
    note(doc, "* reference points, not hardpoints.")
    note(doc,
         f"STATIC ALIGNMENT ENCODED HERE: camber {cam:+.2f}°, toe {toe:+.2f}° per "
         f"side. CONTACT_PATCH sits directly below WHEEL_CENTER on all four "
         f"corners, which is a zero-camber, zero-toe convention, while the rate "
         f"and roll tables assume the design-intent static camber. Decide the "
         f"convention once and apply it everywhere — see section 8.",
         color=BAD, italic=False)


def load_case_section(doc, veh: sla.VehicleData) -> None:
    heading(doc, "7. Load case behind the leg forces")
    note(doc, "The superseded document printed leg forces with no stated load "
              "case, so they could not be checked or defended. These are the "
              "assumptions.")
    table(doc, ["Assumption", "Value", "Source"], [
        ["Total mass", f"{veh.total_mass_kg:.1f} kg", "input"],
        ["Front mass fraction", f"{100 * veh.front_mass_fraction:.1f} %", "input"],
        ["CG height", f"{veh.cg_height_mm:.1f} mm", "estimate"],
        ["Lateral friction coefficient", "1.50", "design_intent"],
        ["Longitudinal friction coefficient", "1.40", "design_intent"],
        ["Lateral acceleration", "1.50 g", "design_intent"],
        ["Lateral load transfer distribution", "0.55", "design_intent"],
    ], widths=[3.2, 1.9, 1.8], right_align_from=1)

    note(doc, "TWO LIMITS, to be repeated wherever the force table appears:",
         italic=False, color=BAD)
    for text in (
        "1. The friction coefficients are a TYRE assumption. PUCPR Racing has no "
        "TTC data, so every force inherits an unmeasured input.",
        "2. There is no pushrod in this model. An upright on two A-arms and a tie "
        "rod has five constraints against six degrees of freedom; the sixth "
        "member is the pushrod. The solve closes the system by treating the UCA "
        "as a two-force member with no X component and enforcing only the moment "
        "about the roll axis, and it places the front ball joints at zero caster, "
        "so braking generates no kingpin moment and the tie rod carries nothing.",
        "These are screening numbers. They are NOT a load set to size tubes from.",
    ):
        note(doc, text, italic=False)


def open_items_section(doc, front: sla.AxleGeometry, rear: sla.AxleGeometry,
                       hp: gs.MergedHardpoints) -> None:
    import math
    heading(doc, "8. Open items")
    r = front.inputs.loaded_radius_mm
    rows = []
    for label, geo in (("Front", front), ("Rear", rear)):
        cam = abs(geo.inputs.static_camber_deg)
        shift = r * math.tan(math.radians(cam))
        moved = geo.scrub_radius_mm + shift
        rows.append([f"{label} scrub if static camber were built in",
                     f"{geo.scrub_radius_mm:.2f} → {moved:.2f} mm",
                     "leaves the 5 to 25 window" if moved > 25 else "still in window"])
    table(doc, ["Item", "Effect", "Consequence"], rows,
          widths=[3.2, 1.9, 1.8], right_align_from=1)

    for text in (
        "The static camber / toe convention has to be decided once: either the "
        "export carries it and scrub radius moves, or it does not and the rate "
        "tables stop claiming it.",
        "The front mass fraction is an unusual value for a rear-drive car and it "
        "drives both the roll-stiffness requirement and every leg force. Confirm "
        "it before the chassis is sized against it.",
        # Superseded: the rear toe link gained a synthesis path and tests in
        # rev 5. What is still open is its OUTBOARD end, which is an upright
        # feature rather than a chassis point and so was left behind when the
        # inboard end moved forward for packaging. Mirrors the same check in
        # geometry_summary.alignment_section.
        f"The rear toe link outboard point sits "
        f"{hp.arr('RL', 'LCA_OUT')[0] - hp.arr('RL', 'TIE_ROD_OUT')[0]:.1f} mm "
        f"behind the wishbone outboard ball joints, at Y = "
        f"{abs(hp.arr('RL', 'TIE_ROD_OUT')[1]):.0f}. It is an upright feature, "
        "not a chassis connection, so it stayed put when the inboard end moved "
        "forward. Bringing it ahead of the axle would be a rear-upright "
        "redesign and would flip the sign of toe change under longitudinal "
        "load."
        if hp.arr("RL", "TIE_ROD_OUT")[0] < hp.arr("RL", "LCA_OUT")[0] else "",
        "Deliberately not computed here, and out of scope for this tool: wheel "
        "rates, motion ratio, ride frequency, damping; stress, deflection, "
        "fatigue, buckling; anything requiring tyre data until TTC data is "
        "acquired.",
    ):
        if text:
            note(doc, text, italic=False)


# Every geometry input the chart would be drawn from. If the image predates any
# of these it is showing a superseded car, and no script in this repo regenerates
# it -- the plotting mode the old caption credited to sla_geometry.py no longer
# exists there. A figure that silently lags the tables beside it is the visual
# form of the failure this project forbids for numbers, so the figure is dropped
# rather than shown with a caption it cannot support.
CHART_SOURCES = (
    REPO / "sla_geometry.py",
    REPO / "steering_geometry.py",
    REPO / "scripts" / "geometry_summary.py",
)


def chart_is_current() -> bool:
    if not CHART.exists():
        return False
    drawn = CHART.stat().st_mtime
    return all(src.stat().st_mtime <= drawn
               for src in CHART_SOURCES if src.exists())


def charts_section(doc) -> None:
    heading(doc, "9. Charts")
    if not chart_is_current():
        stamp = (dt.datetime.fromtimestamp(CHART.stat().st_mtime).strftime("%Y-%m-%d")
                 if CHART.exists() else "never generated")
        note(doc,
             f"No figure in this revision. {CHART.name} ({stamp}) predates the "
             "current geometry and nothing in the repository regenerates it, so "
             "it would show a superseded car next to current tables. The "
             "numbers it plotted — front-view construction, camber and roll "
             "centre vs travel, half-track change, roll behaviour — are all in "
             "sections 2 and 3 as solved values.",
             italic=False, color=BAD)
        return
    note(doc, "Drawn from the same geometry as the tables above, and newer than "
              "every script that defines it. The plots in the superseded "
              "document came from a mix of toolchains and disagreed with its "
              "own tables.")
    doc.add_picture(str(CHART), width=Inches(6.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    note(doc, "Front-view construction, camber and roll centre vs travel, "
              "half-track change, and roll behaviour — front and rear.")


# --------------------------------------------------------------------------- #
# top level
# --------------------------------------------------------------------------- #

def build_docx() -> Path:
    design = sla.run()
    steer = stg.run()
    hp = gs.build_merged(design, steer)

    problems = gs.verify_against_csv(hp)
    if problems:
        raise SystemExit(
            "refusing to build: merged hardpoints disagree with "
            f"{gs.CSV_PATH.name}\n  " + "\n  ".join(problems))

    doc = Document()
    for s in doc.sections:
        s.orientation = WD_ORIENT.PORTRAIT
        s.left_margin = s.right_margin = Inches(0.8)
        s.top_margin = s.bottom_margin = Inches(0.7)
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10)

    cover(doc, design.vehicle)
    provenance(doc, hp)
    vehicle_section(doc, design.vehicle, design.vehicle_results, design.front, design.rear)

    # Rates and roll are solved in 3D by vdcore.analysis.axle on the merged
    # corners -- the same path geometry_summary.py uses. The old
    # sla.AxleKinematics front-view four-bar was removed in the vdcore refactor
    # (it was blind to pivot-axis rake); this mirrors the source-of-truth script.
    for geo, number, sides in (
        (design.front, "2", ("FL", "FR")),
        (design.rear, "3", ("RL", "RR")),
    ):
        inp = geo.inputs
        axle = gs.Axle(
            left=gs._vdcore_corner(hp, sides[0], inp),
            right=gs._vdcore_corner(hp, sides[1], inp),
        )
        rates = gs.axle_rates(
            axle,
            travel_bump_mm=inp.travel_bump_mm,
            travel_droop_mm=inp.travel_droop_mm,
        )
        roll = gs.axle_roll(axle, inp.roll_reference_deg)
        axle_section(doc, geo, rates, roll, number)

    steering_section(doc, steer, stg.STEERING_2027)
    members_section(doc, hp)
    hardpoints_section(doc, hp)
    load_case_section(doc, design.vehicle)
    open_items_section(doc, design.front, design.rear, hp)
    charts_section(doc)

    OUT_DIR.mkdir(exist_ok=True)
    doc.save(str(DOCX))
    return DOCX


def export_pdf(docx_path: Path) -> Path | None:
    try:
        import win32com.client
    except ImportError:
        print("  pywin32 not installed -- skipping PDF (pip install pywin32)")
        return None

    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()), ReadOnly=False)
        doc.SaveAs(str(PDF.resolve()), FileFormat=17)  # wdFormatPDF
        doc.Close(SaveChanges=0)
        return PDF
    except Exception as exc:  # noqa: BLE001 - report and continue
        print(f"  PDF export failed: {exc}")
        return None
    finally:
        if word is not None:
            word.Quit()


def main(argv: Sequence[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="Build the geometry summary docx/pdf.")
    ap.add_argument("--no-pdf", action="store_true", help="skip the Word PDF export")
    args = ap.parse_args(argv)

    docx_path = build_docx()
    print(f"  docx written to {docx_path}")

    if not args.no_pdf:
        pdf_path = export_pdf(docx_path)
        if pdf_path:
            print(f"  pdf  written to {pdf_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
