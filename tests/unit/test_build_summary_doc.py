"""The published .docx/.pdf must still build, and say what the scripts say.

WHY THIS FILE EXISTS
    `scripts/build_summary_doc.py` produces the two artefacts the team and the
    Design judges actually read. Nothing imported it, so nothing noticed when
    rev 5 removed `MergedHardpoints.rear_tie_rod_from_csv` and left two live
    references to it behind: the script raised `AttributeError` on every run
    from 2026-09-01 onward. The .docx and .pdf in `Geometry Summary/` therefore
    sat two geometry revisions out of date while the whole suite stayed green
    and the Markdown summary regenerated cleanly -- the failure was invisible
    from inside the tests.

WHAT THIS PINS
    Not values -- those belong in the benchmarks. This pins that the document
    BUILDS, that it carries the current geometry rather than a cached copy, and
    that the provenance disclosures it makes are ones the code can still back.
"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

REPO = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO))
sys.path.insert(0, str(REPO / "scripts"))

pytest.importorskip("docx", reason="python-docx is not a declared dependency")

import build_summary_doc as bsd  # noqa: E402
import geometry_summary as gs  # noqa: E402

import sla_geometry as sla  # noqa: E402


@pytest.fixture(scope="module")
def built(tmp_path_factory: pytest.TempPathFactory) -> str:
    """Build the .docx into a temp dir and return all of its text.

    The PDF leg is not exercised: it drives Word over COM and is not available
    everywhere. `build_docx` is the part that broke.
    """
    from docx import Document

    out = tmp_path_factory.mktemp("summary")
    docx = out / "summary.docx"
    bsd.OUT_DIR, bsd.DOCX = out, docx
    bsd.build_docx()

    doc = Document(str(docx))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def test_document_builds(built: str) -> None:
    """The regression itself: it raised AttributeError for a day."""
    assert len(built) > 5_000


def test_carries_the_current_rear_geometry(built: str) -> None:
    """Rev-6 hardpoints, not a cached rev-5 copy.

    The LCA pickups are the points rev 6 moved; if the document were rebuilt
    from anything but the live scripts these would still read -1300/-1460.
    """
    rl = gs.build_merged(sla.run(), gs.stg.run()).points["RL"]
    for name in ("LCA_IN_FRONT", "LCA_IN_REAR", "TIE_ROD_IN"):
        assert f"{rl[name][0]:.2f}" in built, f"{name} is stale in the document"


def test_member_band_matches_the_scripts(built: str) -> None:
    """Section 5's band is the one `sla.CheckLimits` actually enforces.

    Rev 6 raised the sla band to 490 and left `ARM_LENGTH_WINDOW_MM` at 460, so
    the generated Markdown failed the rear LCA front leg in one section and
    passed the identical number in another. The .docx is built from the same
    constant and would have inherited it.
    """
    lo, hi = gs.ARM_LENGTH_WINDOW_MM
    assert (lo, hi) == sla.CheckLimits().lca_length_mm
    assert f"{lo:.0f} to {hi:.0f}" in built
    longest = max(sla.member_legs_mm(sla.run().rear).values())
    assert lo <= longest <= hi


def test_no_superseded_provenance_claims(built: str) -> None:
    """The rear toe link stopped being a CSV read-back in rev 5.

    The document used to print, in red, that RL/RR TIE_ROD_* were hand-entered
    points read out of `carro_formula_2027.csv` and covered by no test. That is
    no longer true, and a false open item is as costly as a missing one.
    """
    for stale in ("hand-entered", "no synthesis script",
                  "not covered by any test"):
        assert stale not in built, f"superseded claim still printed: {stale!r}"
    # ...but the design_intent provenance must still be disclosed.
    assert gs.REAR_TOE_LINK_SOURCE in built


def test_stale_chart_is_not_embedded_silently(built: str) -> None:
    """A figure that lags the tables is the visual form of a silent wrong number.

    `geometria.png` predates rev 5 and no script in the repo regenerates it, so
    the charts section must say so rather than print it under a caption
    claiming it was regenerated.
    """
    if bsd.chart_is_current():
        pytest.skip("chart is current; the staleness path is not exercised")
    assert "No figure in this revision" in built
    assert "Regenerated from sla_geometry.py" not in built
